import os
import re
import io
import json
import math
import time
import uuid
import shutil
import asyncio
import zipfile
import tempfile

from pathlib import Path
from typing import Any

import httpx
import pandas as pd

from fastapi import (
    FastAPI,
    UploadFile,
    File,
    Form,
    HTTPException,
)

from fastapi.responses import (
    JSONResponse,
    FileResponse,
)

from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

NODE_ROLE = os.getenv(
    "NODE_ROLE",
    "coordinator"
).lower()

NODE_NAME = os.getenv(
    "NODE_NAME",
    "NODE-1"
)

PORT = int(
    os.getenv(
        "PORT",
        "8000"
    )
)

# ------------------------------------------------------------
# Coordinator URL
# ------------------------------------------------------------

COORDINATOR_URL = os.getenv(
    "COORDINATOR_URL",
    ""
).rstrip("/")

# ------------------------------------------------------------
# Worker URLs
#
# Example:
#
# WORKER_URLS=https://node2.onrender.com,https://node3.onrender.com
# ------------------------------------------------------------

WORKER_URLS = [
    url.strip().rstrip("/")
    for url in os.getenv(
        "WORKER_URLS",
        ""
    ).split(",")
    if url.strip()
]

# ------------------------------------------------------------
# Security
# ------------------------------------------------------------

PROGRESS_SECRET = os.getenv(
    "PROGRESS_SECRET",
    ""
)

# ------------------------------------------------------------
# Distribution model
#
# Your current measured rate:
#
# 0.3 files/sec
#
# 0.3 * 180 seconds = 54 files
# ------------------------------------------------------------

FILES_PER_SECOND = float(
    os.getenv(
        "FILES_PER_SECOND",
        "0.3"
    )
)

TARGET_SECONDS = float(
    os.getenv(
        "TARGET_SECONDS",
        "180"
    )
)

CAPACITY_PER_NODE = int(
    os.getenv(
        "CAPACITY_PER_NODE",
        str(
            math.floor(
                FILES_PER_SECOND
                *
                TARGET_SECONDS
            )
        )
    )
)

# ------------------------------------------------------------
# Storage
# ------------------------------------------------------------

BASE_DIR = Path(
    os.getenv(
        "AUTOGEN_DATA_DIR",
        "./autogen_data"
    )
)

BASE_DIR.mkdir(
    parents=True,
    exist_ok=True
)


# ============================================================
# FASTAPI
# ============================================================

app = FastAPI(
    title="Autogen Distributed Receipt Generator",
    version="2.0.0"
)


# ============================================================
# IN-MEMORY JOB STORE
# ============================================================

jobs: dict[str, dict[str, Any]] = {}


# ============================================================
# LOGGING
# ============================================================

def log(message: str):

    timestamp = time.strftime(
        "%H:%M:%S"
    )

    print(
        f"[{timestamp}] {message}",
        flush=True
    )


def log_line():

    print(
        "-" * 90,
        flush=True
    )


# ============================================================
# VALUE FORMATTER
# ============================================================

def format_value(value):

    if pd.isna(value):
        return ""

    if isinstance(
        value,
        pd.Timestamp
    ):

        return value.strftime(
            "%d/%m/%Y"
        )

    if (
        isinstance(value, float)
        and value.is_integer()
    ):

        return str(
            int(value)
        )

    return str(value)


# ============================================================
# SAFE FILE NAME
# ============================================================

def safe_filename(value):

    value = str(
        value
    ).strip()

    value = re.sub(
        r'[<>:"/\\|?*]',
        "_",
        value
    )

    return value


# ============================================================
# PLACEHOLDER REPLACEMENT
# ============================================================

def replace_placeholders_in_paragraph(
    paragraph,
    data
):

    runs = paragraph.runs

    if not runs:
        return

    pattern = r"\{\{([^{}]+)\}\}"

    # --------------------------------------------------------
    # Normal replacements
    # --------------------------------------------------------

    for run in runs:

        if not run.text:
            continue

        def replace_match(match):

            field = (
                match.group(1)
                .strip()
            )

            if field in data:

                return format_value(
                    data[field]
                )

            return match.group(0)

        run.text = re.sub(
            pattern,
            replace_match,
            run.text
        )

    # --------------------------------------------------------
    # Rebuild paragraph
    # --------------------------------------------------------

    runs = paragraph.runs

    full_text = "".join(
        run.text or ""
        for run in runs
    )

    if "{{" not in full_text:
        return

    matches = list(
        re.finditer(
            pattern,
            full_text
        )
    )

    if not matches:
        return

    # --------------------------------------------------------
    # Run positions
    # --------------------------------------------------------

    run_positions = []

    position = 0

    for index, run in enumerate(runs):

        text = run.text or ""

        start = position

        end = (
            position
            +
            len(text)
        )

        run_positions.append({
            "index": index,
            "start": start,
            "end": end
        })

        position = end

    # --------------------------------------------------------
    # Right to left
    # --------------------------------------------------------

    for match in reversed(matches):

        field = (
            match.group(1)
            .strip()
        )

        if field not in data:
            continue

        replacement = format_value(
            data[field]
        )

        placeholder_start = (
            match.start()
        )

        placeholder_end = (
            match.end()
        )

        affected = []

        for item in run_positions:

            if (
                item["end"] > placeholder_start
                and
                item["start"] < placeholder_end
            ):

                affected.append(
                    item
                )

        if not affected:
            continue

        first = affected[0]

        last = affected[-1]

        first_run = runs[
            first["index"]
        ]

        last_run = runs[
            last["index"]
        ]

        # ----------------------------------------------------
        # Same run
        # ----------------------------------------------------

        if (
            first["index"]
            ==
            last["index"]
        ):

            text = first_run.text

            local_start = (
                placeholder_start
                -
                first["start"]
            )

            local_end = (
                placeholder_end
                -
                first["start"]
            )

            first_run.text = (
                text[:local_start]
                +
                replacement
                +
                text[local_end:]
            )

        else:

            # ------------------------------------------------
            # First run
            # ------------------------------------------------

            first_text = (
                first_run.text
            )

            local_start = (
                placeholder_start
                -
                first["start"]
            )

            first_run.text = (
                first_text[:local_start]
                +
                replacement
            )

            # ------------------------------------------------
            # Middle runs
            # ------------------------------------------------

            for item in affected[1:-1]:

                runs[
                    item["index"]
                ].text = ""

            # ------------------------------------------------
            # Last run
            # ------------------------------------------------

            last_text = (
                last_run.text
            )

            local_end = (
                placeholder_end
                -
                last["start"]
            )

            last_run.text = (
                last_text[local_end:]
            )


# ============================================================
# TABLE
# ============================================================

def process_table(
    table,
    data
):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_placeholders_in_paragraph(
                    paragraph,
                    data
                )

            for nested_table in cell.tables:

                process_table(
                    nested_table,
                    data
                )


# ============================================================
# HEADER / FOOTER
# ============================================================

def process_header_footer(
    header_footer,
    data
):

    for paragraph in (
        header_footer.paragraphs
    ):

        replace_placeholders_in_paragraph(
            paragraph,
            data
        )

    for table in (
        header_footer.tables
    ):

        process_table(
            table,
            data
        )


# ============================================================
# COMPLETE DOCUMENT
# ============================================================

def process_document(
    document,
    data
):

    for paragraph in (
        document.paragraphs
    ):

        replace_placeholders_in_paragraph(
            paragraph,
            data
        )

    for table in document.tables:

        process_table(
            table,
            data
        )

    for section in document.sections:

        process_header_footer(
            section.header,
            data
        )

        process_header_footer(
            section.footer,
            data
        )


# ============================================================
# FIND PLACEHOLDERS
# ============================================================

def get_placeholders(
    document
):

    placeholders = set()

    pattern = r"\{\{([^{}]+)\}\}"

    def scan_text(text):

        matches = re.findall(
            pattern,
            text
        )

        for match in matches:

            placeholders.add(
                match.strip()
            )

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = "".join(
            run.text or ""
            for run in paragraph.runs
        )

        scan_text(text)

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    def scan_table(table):

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    text = "".join(
                        run.text or ""
                        for run in paragraph.runs
                    )

                    scan_text(text)

                for nested in cell.tables:

                    scan_table(
                        nested
                    )

    for table in document.tables:

        scan_table(table)

    # --------------------------------------------------------
    # Headers / Footers
    # --------------------------------------------------------

    for section in document.sections:

        for header_footer in [
            section.header,
            section.footer
        ]:

            for paragraph in (
                header_footer.paragraphs
            ):

                text = "".join(
                    run.text or ""
                    for run in paragraph.runs
                )

                scan_text(text)

            for table in (
                header_footer.tables
            ):

                scan_table(table)

    return sorted(
        placeholders
    )


# ============================================================
# GENERATE ONE RECEIPT
# ============================================================

def generate_receipt(
    template_file,
    output_folder,
    data,
    row_number
):

    started = time.perf_counter()

    document = Document(
        str(template_file)
    )

    process_document(
        document,
        data
    )

    name = format_value(
        data.get(
            "name",
            f"Client_{row_number}"
        )
    )

    account = format_value(
        data.get(
            "acc_no",
            ""
        )
    )

    name = safe_filename(
        name
    )

    account = safe_filename(
        account
    )

    if not name:

        name = (
            f"Client_{row_number}"
        )

    if account:

        filename = (
            f"{name}_{account}.docx"
        )

    else:

        filename = (
            f"{name}.docx"
        )

    output_path = (
        Path(output_folder)
        /
        filename
    )

    original_path = output_path

    counter = 1

    while output_path.exists():

        output_path = (
            original_path.parent
            /
            f"{original_path.stem}_"
            f"{counter}"
            f"{original_path.suffix}"
        )

        counter += 1

    document.save(
        str(output_path)
    )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    return (
        output_path,
        elapsed
    )


# ============================================================
# CREATE ZIP
# ============================================================

def create_zip(
    source_folder,
    zip_path
):

    started = time.perf_counter()

    source_folder = Path(
        source_folder
    )

    zip_path = Path(
        zip_path
    )

    file_count = 0

    with zipfile.ZipFile(
        zip_path,
        "w",
        compression=zipfile.ZIP_DEFLATED
    ) as archive:

        for file in source_folder.rglob("*"):

            if not file.is_file():
                continue

            archive.write(
                file,
                arcname=file.relative_to(
                    source_folder
                )
            )

            file_count += 1

    elapsed = (
        time.perf_counter()
        -
        started
    )

    return (
        zip_path,
        file_count,
        elapsed
    )


# ============================================================
# SPLIT ROWS
# ============================================================

def split_rows(
    rows,
    node_count
):

    if node_count <= 1:

        return [
            rows
        ]

    total = len(rows)

    base = (
        total
        //
        node_count
    )

    remainder = (
        total
        %
        node_count
    )

    chunks = []

    start = 0

    for index in range(
        node_count
    ):

        size = (
            base
            +
            (
                1
                if index < remainder
                else 0
            )
        )

        chunks.append(
            rows[
                start:
                start + size
            ]
        )

        start += size

    return chunks


# ============================================================
# NODE COUNT
# ============================================================

def calculate_node_count(
    total_files
):

    available_nodes = (
        1
        +
        len(WORKER_URLS)
    )

    required_nodes = max(
        1,
        math.ceil(
            total_files
            /
            CAPACITY_PER_NODE
        )
    )

    return min(
        required_nodes,
        available_nodes
    )


# ============================================================
# UPDATE GLOBAL PROGRESS
# ============================================================

def recalculate_job_progress(
    job
):

    nodes = job.get(
        "nodes",
        {}
    )

    completed = sum(
        int(
            node.get(
                "completed",
                0
            )
        )
        for node in nodes.values()
    )

    total = job.get(
        "total",
        0
    )

    if total:

        progress = int(
            completed
            /
            total
            *
            100
        )

    else:

        progress = 0

    job["completed"] = completed

    job["progress"] = progress


# ============================================================
# PROGRESS CALLBACK
# ============================================================

async def send_progress_to_coordinator(

    coordinator_url,
    job_id,
    node_name,
    completed,
    total,
    elapsed,
    rate,
    filename,
    file_time=None,
    status="running"

):

    if not coordinator_url:

        return

    payload = {

        "secret":
            PROGRESS_SECRET,

        "job_id":
            job_id,

        "node_name":
            node_name,

        "completed":
            completed,

        "total":
            total,

        "elapsed":
            elapsed,

        "rate":
            rate,

        "filename":
            filename,

        "file_time":
            file_time,

        "status":
            status
    }

    try:

        timeout = httpx.Timeout(
            10.0,
            connect=5.0
        )

        async with httpx.AsyncClient(
            timeout=timeout
        ) as client:

            response = await client.post(

                f"{coordinator_url}"
                "/internal/progress",

                json=payload
            )

            if response.status_code != 200:

                log(
                    f"{node_name}: "
                    f"progress callback returned "
                    f"{response.status_code}"
                )

    except Exception as error:

        log(
            f"{node_name}: "
            f"progress callback failed: "
            f"{error}"
        )


# ============================================================
# WAKE ONE WORKER
# ============================================================

async def wake_worker(
    url,
    worker_number
):

    worker_name = (
        f"NODE-{worker_number}"
    )

    health_url = (
        f"{url}/health"
    )

    log(
        f"{worker_name}: "
        f"waking {health_url}"
    )

    started = time.perf_counter()

    timeout = httpx.Timeout(
        60.0,
        connect=60.0
    )

    for attempt in range(
        1,
        11
    ):

        try:

            async with httpx.AsyncClient(
                timeout=timeout
            ) as client:

                response = await client.get(
                    health_url
                )

                if (
                    response.status_code
                    ==
                    200
                ):

                    elapsed = (
                        time.perf_counter()
                        -
                        started
                    )

                    log(
                        f"{worker_name}: "
                        f"READY "
                        f"in {elapsed:.2f}s"
                    )

                    return {
                        "node":
                            worker_name,

                        "url":
                            url,

                        "ready":
                            True,

                        "wake_time":
                            elapsed
                    }

        except Exception as error:

            log(
                f"{worker_name}: "
                f"attempt {attempt}/10 "
                f"not ready"
            )

        await asyncio.sleep(
            2
        )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    log(
        f"{worker_name}: "
        f"FAILED TO WAKE "
        f"after {elapsed:.2f}s"
    )

    return {
        "node":
            worker_name,

        "url":
            url,

        "ready":
            False,

        "wake_time":
            elapsed
    }


# ============================================================
# WAKE ALL WORKERS
# ============================================================

async def wake_all_workers():

    if not WORKER_URLS:

        log(
            "No workers configured."
        )

        return []

    log_line()

    log(
        f"WAKING "
        f"{len(WORKER_URLS)} "
        f"WORKER(S)"
    )

    started = time.perf_counter()

    tasks = []

    for index, url in enumerate(
        WORKER_URLS,
        start=2
    ):

        tasks.append(
            wake_worker(
                url,
                index
            )
        )

    results = await asyncio.gather(
        *tasks
    )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    ready = sum(
        1
        for result in results
        if result["ready"]
    )

    log(
        f"Workers ready: "
        f"{ready}/{len(results)}"
    )

    log(
        f"Wake phase: "
        f"{elapsed:.2f}s"
    )

    log_line()

    return results


# ============================================================
# LOCAL NODE PROCESSING
# ============================================================

async def process_local_chunk(

    job_id,
    template_bytes,
    rows,
    node_name

):

    job = jobs[job_id]

    node = job[
        "nodes"
    ][node_name]

    node["status"] = "running"

    node_dir = (
        BASE_DIR
        /
        job_id
        /
        node_name
    )

    node_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    template_path = (
        node_dir
        /
        "template.docx"
    )

    template_path.write_bytes(
        template_bytes
    )

    output_dir = (
        node_dir
        /
        "receipts"
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    total = len(rows)

    completed = 0

    errors = []

    started = time.perf_counter()

    log(
        f"[{node_name}] "
        f"START "
        f"{total} files"
    )

    for index, row in enumerate(
        rows
    ):

        file_number = (
            index + 1
        )

        file_started = (
            time.perf_counter()
        )

        try:

            output_file, file_elapsed = (
                await asyncio.to_thread(

                    generate_receipt,

                    template_path,

                    output_dir,

                    row,

                    file_number
                )
            )

            completed += 1

            elapsed = (
                time.perf_counter()
                -
                started
            )

            rate = (
                completed
                /
                elapsed
                if elapsed > 0
                else 0
            )

            remaining = (
                total
                -
                completed
            )

            eta = (
                remaining
                /
                rate
                if rate > 0
                else 0
            )

            progress = (
                int(
                    completed
                    /
                    total
                    *
                    100
                )
                if total
                else 100
            )

            node[
                "completed"
            ] = completed

            node[
                "progress"
            ] = progress

            node[
                "elapsed"
            ] = elapsed

            node[
                "rate"
            ] = rate

            node[
                "eta"
            ] = eta

            node[
                "last_file"
            ] = output_file.name

            node[
                "last_file_time"
            ] = file_elapsed

            recalculate_job_progress(
                job
            )

            log(
                f"[{node_name}] "
                f"{file_number}/{total} "
                f"| {output_file.name} "
                f"| file={file_elapsed:.3f}s "
                f"| total={elapsed:.2f}s "
                f"| rate={rate:.3f}/s "
                f"| ETA={eta:.2f}s "
                f"| {progress}% "
                f"| GLOBAL="
                f"{job['completed']}/"
                f"{job['total']} "
                f"({job['progress']}%)"
            )

        except Exception as error:

            file_elapsed = (
                time.perf_counter()
                -
                file_started
            )

            errors.append({
                "row":
                    file_number,

                "error":
                    str(error),

                "time":
                    file_elapsed
            })

            log(
                f"[{node_name}] "
                f"ERROR "
                f"{file_number}/{total} "
                f"| time="
                f"{file_elapsed:.3f}s "
                f"| {error}"
            )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    rate = (
        completed
        /
        elapsed
        if elapsed > 0
        else 0
    )

    node[
        "completed"
    ] = completed

    node[
        "progress"
    ] = 100

    node[
        "elapsed"
    ] = elapsed

    node[
        "rate"
    ] = rate

    node[
        "eta"
    ] = 0

    node[
        "errors"
    ] = errors

    node[
        "status"
    ] = "complete"

    recalculate_job_progress(
        job
    )

    log(
        f"[{node_name}] "
        f"COMPLETE "
        f"{completed}/{total} "
        f"| elapsed={elapsed:.3f}s "
        f"| rate={rate:.3f}/s "
        f"| errors={len(errors)}"
    )

    return {
        "node":
            node_name,

        "directory":
            str(output_dir),

        "completed":
            completed,

        "total":
            total,

        "elapsed":
            elapsed,

        "rate":
            rate,

        "errors":
            errors
    }


# ============================================================
# WORKER PROCESSING
# ============================================================

async def process_worker_chunk(

    job_id,
    template_bytes,
    rows,
    node_name,
    coordinator_url

):

    node_dir = (
        BASE_DIR
        /
        job_id
        /
        node_name
    )

    node_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    template_path = (
        node_dir
        /
        "template.docx"
    )

    template_path.write_bytes(
        template_bytes
    )

    output_dir = (
        node_dir
        /
        "receipts"
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    total = len(rows)

    completed = 0

    errors = []

    started = time.perf_counter()

    log(
        f"[{node_name}] "
        f"START "
        f"{total} files"
    )

    progress_tasks = []

    for index, row in enumerate(
        rows
    ):

        file_number = (
            index + 1
        )

        try:

            output_file, file_elapsed = (
                await asyncio.to_thread(

                    generate_receipt,

                    template_path,

                    output_dir,

                    row,

                    file_number
                )
            )

            completed += 1

            elapsed = (
                time.perf_counter()
                -
                started
            )

            rate = (
                completed
                /
                elapsed
                if elapsed > 0
                else 0
            )

            remaining = (
                total
                -
                completed
            )

            eta = (
                remaining
                /
                rate
                if rate > 0
                else 0
            )

            progress = (
                int(
                    completed
                    /
                    total
                    *
                    100
                )
                if total
                else 100
            )

            log(
                f"[{node_name}] "
                f"{file_number}/{total} "
                f"| {output_file.name} "
                f"| file="
                f"{file_elapsed:.3f}s "
                f"| total="
                f"{elapsed:.2f}s "
                f"| rate="
                f"{rate:.3f}/s "
                f"| ETA="
                f"{eta:.2f}s "
                f"| {progress}%"
            )

            # ------------------------------------------------
            # Send progress without blocking generation
            # ------------------------------------------------

            progress_tasks.append(

                asyncio.create_task(

                    send_progress_to_coordinator(

                        coordinator_url,

                        job_id,

                        node_name,

                        completed,

                        total,

                        elapsed,

                        rate,

                        output_file.name,

                        file_elapsed,

                        "running"
                    )
                )
            )

            # Keep memory bounded.
            if len(progress_tasks) >= 20:

                done, pending = (
                    await asyncio.wait(
                        progress_tasks,
                        return_when=
                        asyncio.FIRST_COMPLETED
                    )
                )

                progress_tasks = list(
                    pending
                )

        except Exception as error:

            errors.append({
                "row":
                    file_number,

                "error":
                    str(error)
            })

            log(
                f"[{node_name}] "
                f"ERROR "
                f"{file_number}/{total} "
                f"| {error}"
            )

    # --------------------------------------------------------
    # Wait for outstanding progress calls
    # --------------------------------------------------------

    if progress_tasks:

        await asyncio.gather(
            *progress_tasks,
            return_exceptions=True
        )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    rate = (
        completed
        /
        elapsed
        if elapsed > 0
        else 0
    )

    # --------------------------------------------------------
    # Final progress callback
    # --------------------------------------------------------

    await send_progress_to_coordinator(

        coordinator_url,

        job_id,

        node_name,

        completed,

        total,

        elapsed,

        rate,

        "",

        None,

        "complete"
    )

    log(
        f"[{node_name}] "
        f"COMPLETE "
        f"{completed}/{total} "
        f"| elapsed="
        f"{elapsed:.3f}s "
        f"| rate="
        f"{rate:.3f}/s "
        f"| errors="
        f"{len(errors)}"
    )

    # --------------------------------------------------------
    # Create worker ZIP
    # --------------------------------------------------------

    zip_started = (
        time.perf_counter()
    )

    zip_path = (
        node_dir
        /
        f"{node_name}.zip"
    )

    _, zip_file_count, zip_elapsed = (
        create_zip(
            output_dir,
            zip_path
        )
    )

    log(
        f"[{node_name}] "
        f"ZIP CREATED "
        f"| files={zip_file_count} "
        f"| time={zip_elapsed:.3f}s"
    )

    total_time = (
        time.perf_counter()
        -
        started
    )

    return {
        "node":
            node_name,

        "zip_path":
            str(zip_path),

        "completed":
            completed,

        "total":
            total,

        "elapsed":
            elapsed,

        "rate":
            rate,

        "errors":
            errors,

        "zip_elapsed":
            zip_elapsed,

        "total_time":
            total_time
    }


# ============================================================
# SEND CHUNK TO WORKER
# ============================================================

async def send_chunk_to_worker(

    worker_url,
    template_bytes,
    rows,
    job_id,
    node_name

):

    started = time.perf_counter()

    log(
        f"[{node_name}] "
        f"SENDING "
        f"{len(rows)} files "
        f"to {worker_url}"
    )

    files = {

        "template": (
            "template.docx",
            template_bytes,
            "application/"
            "vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    }

    data = {

        "job_id":
            job_id,

        "node_name":
            node_name,

        "coordinator_url":
            COORDINATOR_URL,

        "data_json":
            json.dumps(
                rows,
                default=str
            )
    }

    timeout = httpx.Timeout(
        600.0,
        connect=60.0,
        read=600.0,
        write=600.0
    )

    try:

        async with httpx.AsyncClient(
            timeout=timeout
        ) as client:

            response = await client.post(

                f"{worker_url}/worker/process",

                files=files,

                data=data
            )

        if response.status_code != 200:

            raise RuntimeError(
                f"Worker returned "
                f"{response.status_code}: "
                f"{response.text}"
            )

        # ----------------------------------------------------
        # Save returned ZIP
        # ----------------------------------------------------

        result_dir = (
            BASE_DIR
            /
            job_id
            /
            "worker_results"
        )

        result_dir.mkdir(
            parents=True,
            exist_ok=True
        )

        zip_path = (
            result_dir
            /
            f"{node_name}.zip"
        )

        zip_path.write_bytes(
            response.content
        )

        elapsed = (
            time.perf_counter()
            -
            started
        )

        result = response.headers.get(
            "X-Worker-Result"
        )

        log(
            f"[{node_name}] "
            f"RECEIVED ZIP "
            f"| size="
            f"{len(response.content):,} bytes "
            f"| transfer="
            f"{elapsed:.3f}s"
        )

        return {
            "node":
                node_name,

            "zip_path":
                str(zip_path),

            "transfer_elapsed":
                elapsed,

            "completed":
                len(rows),

            "total":
                len(rows),

            "status":
                "complete",

            "worker_result":
                result
        }

    except Exception as error:

        log(
            f"[{node_name}] "
            f"WORKER FAILED: "
            f"{error}"
        )

        raise


# ============================================================
# COMBINE ALL RECEIPTS
# ============================================================

def combine_zip_files(
    job_id,
    zip_paths
):

    started = time.perf_counter()

    final_dir = (
        BASE_DIR
        /
        job_id
        /
        "final"
    )

    final_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    final_zip = (
        final_dir
        /
        "autogen_receipts.zip"
    )

    file_count = 0

    used_names = set()

    with zipfile.ZipFile(
        final_zip,
        "w",
        compression=zipfile.ZIP_DEFLATED
    ) as final_archive:

        for zip_path in zip_paths:

            if not zip_path:
                continue

            zip_path = Path(
                zip_path
            )

            if not zip_path.exists():

                log(
                    f"ZIP missing: "
                    f"{zip_path}"
                )

                continue

            with zipfile.ZipFile(
                zip_path,
                "r"
            ) as source_zip:

                for info in (
                    source_zip.infolist()
                ):

                    if info.is_dir():
                        continue

                    filename = Path(
                        info.filename
                    ).name

                    if not filename:
                        continue

                    # ------------------------------------------------
                    # Prevent duplicate names
                    # ------------------------------------------------

                    original = filename

                    counter = 1

                    while filename in used_names:

                        stem = Path(
                            original
                        ).stem

                        suffix = Path(
                            original
                        ).suffix

                        filename = (
                            f"{stem}_{counter}"
                            f"{suffix}"
                        )

                        counter += 1

                    used_names.add(
                        filename
                    )

                    data = source_zip.read(
                        info
                    )

                    final_archive.writestr(
                        filename,
                        data
                    )

                    file_count += 1

    elapsed = (
        time.perf_counter()
        -
        started
    )

    return (
        final_zip,
        file_count,
        elapsed
    )


# ============================================================
# HEALTH
# ============================================================

@app.get("/health")
async def health():

    return {
        "status":
            "ok",

        "node":
            NODE_NAME,

        "role":
            NODE_ROLE,

        "workers":
            len(WORKER_URLS)
            if NODE_ROLE
            ==
            "coordinator"
            else 0
    }


# ============================================================
# INSPECT
# ============================================================

@app.post("/inspect")
async def inspect_files(

    template: UploadFile = File(...),

    excel: UploadFile = File(...)

):

    if NODE_ROLE != "coordinator":

        raise HTTPException(
            status_code=403,
            detail="Inspection is only available on the coordinator."
        )

    started = time.perf_counter()

    log_line()

    log(
        "INSPECTION START"
    )

    # --------------------------------------------------------
    # Read template
    # --------------------------------------------------------

    template_bytes = await (
        template.read()
    )

    try:

        template_document = Document(
            io.BytesIO(
                template_bytes
            )
        )

    except Exception as error:

        raise HTTPException(
            status_code=400,
            detail=(
                f"Invalid Word template: "
                f"{error}"
            )
        )

    placeholders = get_placeholders(
        template_document
    )

    log(
        f"Template: "
        f"{template.filename}"
    )

    log(
        f"Fields found: "
        f"{len(placeholders)}"
    )

    # --------------------------------------------------------
    # Read Excel
    # --------------------------------------------------------

    excel_bytes = await (
        excel.read()
    )

    try:

        df = pd.read_excel(
            io.BytesIO(
                excel_bytes
            )
        )

    except Exception as error:

        raise HTTPException(
            status_code=400,
            detail=(
                f"Could not read Excel: "
                f"{error}"
            )
        )

    df.columns = [
        str(column).strip()
        for column in df.columns
    ]

    # --------------------------------------------------------
    # Convert rows
    # --------------------------------------------------------

    rows = []

    for _, row in df.iterrows():

        data = {}

        for column in df.columns:

            value = row[column]

            if pd.isna(value):

                data[column] = ""

            else:

                data[column] = (
                    format_value(value)
                )

        rows.append(
            data
        )

    # --------------------------------------------------------
    # Missing / extra
    # --------------------------------------------------------

    missing_fields = [
        field
        for field in placeholders
        if field not in df.columns
    ]

    extra_columns = [
        column
        for column in df.columns
        if column not in placeholders
    ]

    ready = (
        len(missing_fields)
        ==
        0
        and
        len(rows)
        >
        0
    )

    # --------------------------------------------------------
    # Wake workers
    # --------------------------------------------------------

    wake_results = await (
        wake_all_workers()
    )

    elapsed = (
        time.perf_counter()
        -
        started
    )

    log(
        f"INSPECTION COMPLETE "
        f"in {elapsed:.3f}s"
    )

    log_line()

    return {

        "template": {

            "filename":
                template.filename,

            "fields":
                placeholders
        },

        "excel": {

            "filename":
                excel.filename,

            "columns":
                list(df.columns),

            "row_count":
                len(rows),

            "rows":
                rows
        },

        "validation": {

            "missing_fields":
                missing_fields,

            "extra_columns":
                extra_columns,

            "ready":
                ready
        },

        "settings": {

            "pdf_enabled":
                False
        },

        "workers": {

            "available":
                len(WORKER_URLS) + 1,

            "capacity_per_node":
                CAPACITY_PER_NODE,

            "wake":
                wake_results
        }
    }


# ============================================================
# GENERATE
# ============================================================

@app.post("/generate")
async def generate(

    template: UploadFile = File(...),

    data_json: str = Form(...)

):

    if NODE_ROLE != "coordinator":

        raise HTTPException(
            status_code=403,
            detail="Generation must be requested from the coordinator."
        )

    started = time.perf_counter()

    try:

        rows = json.loads(
            data_json
        )

    except Exception:

        raise HTTPException(
            status_code=400,
            detail="Invalid data_json."
        )

    if not isinstance(
        rows,
        list
    ):

        raise HTTPException(
            status_code=400,
            detail="data_json must be a list."
        )

    if not rows:

        raise HTTPException(
            status_code=400,
            detail="No rows supplied."
        )

    template_bytes = await (
        template.read()
    )

    job_id = uuid.uuid4().hex

    total = len(rows)

    node_count = calculate_node_count(
        total
    )

    chunks = split_rows(
        rows,
        node_count
    )

    # --------------------------------------------------------
    # Job
    # --------------------------------------------------------

    job = {

        "job_id":
            job_id,

        "status":
            "starting",

        "total":
            total,

        "completed":
            0,

        "progress":
            0,

        "failed":
            0,

        "node_count":
            node_count,

        "nodes":
            {},

        "started":
            started,

        "generation_elapsed":
            None,

        "zip_elapsed":
            None,

        "final_elapsed":
            None,

        "final_zip":
            None
    }

    jobs[job_id] = job

    # --------------------------------------------------------
    # Initialize nodes
    # --------------------------------------------------------

    for index, chunk in enumerate(
        chunks
    ):

        node_name = (
            f"NODE-{index + 1}"
        )

        job["nodes"][
            node_name
        ] = {

            "total":
                len(chunk),

            "completed":
                0,

            "progress":
                0,

            "elapsed":
                0,

            "rate":
                0,

            "eta":
                0,

            "status":
                "waiting",

            "last_file":
                None,

            "last_file_time":
                None,

            "errors":
                []
        }

    log_line()

    log(
        "GENERATION START"
    )

    log(
        f"Job: {job_id}"
    )

    log(
        f"Files: {total}"
    )

    log(
        f"Available nodes: "
        f"{len(WORKER_URLS) + 1}"
    )

    log(
        f"Nodes selected: "
        f"{node_count}"
    )

    log(
        f"Capacity/node: "
        f"{CAPACITY_PER_NODE}"
    )

    for index, chunk in enumerate(
        chunks
    ):

        log(
            f"NODE-{index + 1}: "
            f"{len(chunk)} files"
        )

    # --------------------------------------------------------
    # Start background generation
    # --------------------------------------------------------

    asyncio.create_task(

        run_generation(

            job_id,

            template_bytes,

            chunks

        )
    )

    return {

        "job_id":
            job_id,

        "status":
            "started",

        "total":
            total,

        "nodes":
            node_count
    }


# ============================================================
# RUN DISTRIBUTED GENERATION
# ============================================================

async def run_generation(

    job_id,
    template_bytes,
    chunks

):

    job = jobs[job_id]

    generation_started = (
        time.perf_counter()
    )

    job["status"] = (
        "running"
    )

    tasks = []

    # --------------------------------------------------------
    # NODE 1
    # --------------------------------------------------------

    tasks.append(

        process_local_chunk(

            job_id,

            template_bytes,

            chunks[0],

            "NODE-1"
        )
    )

    # --------------------------------------------------------
    # WORKERS
    # --------------------------------------------------------

    for index in range(
        1,
        len(chunks)
    ):

        worker_index = (
            index - 1
        )

        if worker_index >= len(
            WORKER_URLS
        ):

            raise RuntimeError(
                "Not enough worker nodes."
            )

        worker_url = (
            WORKER_URLS[
                worker_index
            ]
        )

        node_name = (
            f"NODE-{index + 1}"
        )

        tasks.append(

            send_chunk_to_worker(

                worker_url,

                template_bytes,

                chunks[index],

                job_id,

                node_name
            )
        )

    # --------------------------------------------------------
    # RUN EVERYTHING CONCURRENTLY
    # --------------------------------------------------------

    results = await asyncio.gather(
        *tasks,
        return_exceptions=True
    )

    generation_elapsed = (
        time.perf_counter()
        -
        generation_started
    )

    job[
        "generation_elapsed"
    ] = generation_elapsed

    log_line()

    log(
        f"DOCUMENT GENERATION PHASE "
        f"COMPLETE "
        f"in {generation_elapsed:.3f}s"
    )

    # --------------------------------------------------------
    # Collect ZIPs
    # --------------------------------------------------------

    zip_paths = []

    for index, result in enumerate(
        results
    ):

        node_name = (
            f"NODE-{index + 1}"
        )

        if isinstance(
            result,
            Exception
        ):

            job["nodes"][
                node_name
            ]["status"] = "error"

            job["nodes"][
                node_name
            ]["error"] = str(
                result
            )

            log(
                f"{node_name}: "
                f"FAILED: "
                f"{result}"
            )

            continue

        job["nodes"][
            node_name
        ].update(
            {
                key: value
                for key, value
                in result.items()
                if key != "zip_path"
            }
        )

        job["nodes"][
            node_name
        ]["status"] = (
            "complete"
        )

        zip_path = result.get(
            "zip_path"
        )

        if zip_path:

            zip_paths.append(
                zip_path
            )

    # --------------------------------------------------------
    # FINAL GLOBAL PROGRESS
    # --------------------------------------------------------

    recalculate_job_progress(
        job
    )

    # --------------------------------------------------------
    # COMBINE ZIP
    # --------------------------------------------------------

    log_line()

    log(
        "CREATING FINAL ZIP..."
    )

    zip_started = (
        time.perf_counter()
    )

    final_zip, file_count, zip_elapsed = (
        await asyncio.to_thread(

            combine_zip_files,

            job_id,

            zip_paths
        )
    )

    job[
        "zip_elapsed"
    ] = zip_elapsed

    log(
        f"FINAL ZIP CREATED "
        f"| files={file_count} "
        f"| time={zip_elapsed:.3f}s "
        f"| size="
        f"{final_zip.stat().st_size:,} bytes"
    )

    # --------------------------------------------------------
    # FINAL TIME
    # --------------------------------------------------------

    final_elapsed = (
        time.perf_counter()
        -
        job["started"]
    )

    job[
        "final_elapsed"
    ] = final_elapsed

    job[
        "final_zip"
    ] = str(
        final_zip
    )

    job[
        "completed"
    ] = min(
        job["completed"],
        job["total"]
    )

    job[
        "progress"
    ] = 100

    job[
        "failed"
    ] = (
        job["total"]
        -
        job["completed"]
    )

    if job["failed"] > 0:

        job[
            "status"
        ] = (
            "completed_with_errors"
        )

    else:

        job[
            "status"
        ] = "completed"

    log_line()

    log(
        "AUTOGEN COMPLETE"
    )

    log(
        f"Job: "
        f"{job_id}"
    )

    log(
        f"Total files: "
        f"{job['total']}"
    )

    log(
        f"Completed: "
        f"{job['completed']}"
    )

    log(
        f"Failed: "
        f"{job['failed']}"
    )

    log(
        f"Generation time: "
        f"{generation_elapsed:.3f}s"
    )

    log(
        f"ZIP creation time: "
        f"{zip_elapsed:.3f}s"
    )

    log(
        f"FINAL TOTAL TIME: "
        f"{final_elapsed:.3f}s"
    )

    log(
        f"Average total/file: "
        f"{final_elapsed / job['total']:.3f}s"
    )

    log(
        f"Final ZIP: "
        f"{final_zip}"
    )

    log_line()


# ============================================================
# PROGRESS CALLBACK
# ============================================================

@app.post("/internal/progress")
async def internal_progress(
    payload: dict
):

    if NODE_ROLE != "coordinator":

        raise HTTPException(
            status_code=403,
            detail="Only coordinator accepts progress."
        )

    if PROGRESS_SECRET:

        if (
            payload.get(
                "secret"
            )
            !=
            PROGRESS_SECRET
        ):

            raise HTTPException(
                status_code=403,
                detail="Invalid progress secret."
            )

    job_id = payload.get(
        "job_id"
    )

    node_name = payload.get(
        "node_name"
    )

    if job_id not in jobs:

        return {
            "accepted":
                False,

            "reason":
                "unknown_job"
        }

    job = jobs[job_id]

    if node_name not in job["nodes"]:

        job["nodes"][
            node_name
        ] = {

            "total":
                payload.get(
                    "total",
                    0
                ),

            "completed":
                0,

            "progress":
                0,

            "elapsed":
                0,

            "rate":
                0,

            "eta":
                0,

            "status":
                "running"
        }

    node = job[
        "nodes"
    ][node_name]

    completed = int(
        payload.get(
            "completed",
            0
        )
    )

    total = int(
        payload.get(
            "total",
            0
        )
    )

    elapsed = float(
        payload.get(
            "elapsed",
            0
        )
    )

    rate = float(
        payload.get(
            "rate",
            0
        )
    )

    file_time = payload.get(
        "file_time"
    )

    filename = payload.get(
        "filename"
    )

    node["total"] = total

    node["completed"] = completed

    node["progress"] = (
        int(
            completed
            /
            total
            *
            100
        )
        if total
        else 100
    )

    node["elapsed"] = elapsed

    node["rate"] = rate

    node["eta"] = (
        (
            total - completed
        )
        /
        rate
        if rate > 0
        else 0
    )

    node["last_file"] = filename

    node["last_file_time"] = (
        file_time
    )

    node["status"] = payload.get(
        "status",
        "running"
    )

    recalculate_job_progress(
        job
    )

    log(
        f"[CALLBACK] "
        f"{node_name} "
        f"{completed}/{total} "
        f"| file={filename or '-'} "
        f"| file_time="
        f"{file_time if file_time is not None else '-'}s "
        f"| GLOBAL="
        f"{job['completed']}/"
        f"{job['total']} "
        f"({job['progress']}%)"
    )

    return {
        "accepted":
            True,

        "global_completed":
            job["completed"],

        "global_total":
            job["total"],

        "global_progress":
            job["progress"]
    }


# ============================================================
# PROGRESS API FOR NEXT.JS
# ============================================================

@app.get("/progress/{job_id}")
async def get_progress(
    job_id: str
):

    job = jobs.get(
        job_id
    )

    if not job:

        raise HTTPException(
            status_code=404,
            detail="Job not found."
        )

    return {

        "job_id":
            job["job_id"],

        "status":
            job["status"],

        "total":
            job["total"],

        "completed":
            job["completed"],

        "failed":
            job["failed"],

        "progress":
            job["progress"],

        "node_count":
            job["node_count"],

        "generation_elapsed":
            job["generation_elapsed"],

        "zip_elapsed":
            job["zip_elapsed"],

        "final_elapsed":
            job["final_elapsed"],

        "nodes":
            job["nodes"],

        "download_url":
            (
                f"/download/{job_id}"
                if job["final_zip"]
                else None
            )
    }


# ============================================================
# DOWNLOAD
# ============================================================

@app.get("/download/{job_id}")
async def download(
    job_id: str
):

    job = jobs.get(
        job_id
    )

    if not job:

        raise HTTPException(
            status_code=404,
            detail="Job not found."
        )

    if not job.get(
        "final_zip"
    ):

        raise HTTPException(
            status_code=409,
            detail="ZIP is not ready."
        )

    zip_path = Path(
        job["final_zip"]
    )

    if not zip_path.exists():

        raise HTTPException(
            status_code=404,
            detail="ZIP file no longer exists."
        )

    return FileResponse(

        path=str(
            zip_path
        ),

        media_type=(
            "application/zip"
        ),

        filename=(
            "autogen_receipts.zip"
        )
    )


# ============================================================
# WORKER ENDPOINT
# ============================================================

@app.post("/worker/process")
async def worker_process(

    template: UploadFile = File(...),

    data_json: str = Form(...),

    job_id: str = Form(...),

    coordinator_url: str = Form(...),

    node_name: str = Form(...)

):

    if NODE_ROLE != "worker":

        raise HTTPException(
            status_code=403,
            detail="This node is not configured as a worker."
        )

    started = time.perf_counter()

    log_line()

    log(
        f"{node_name}: "
        f"JOB RECEIVED "
        f"| job={job_id}"
    )

    try:

        rows = json.loads(
            data_json
        )

    except Exception:

        raise HTTPException(
            status_code=400,
            detail="Invalid data_json."
        )

    if not isinstance(
        rows,
        list
    ):

        raise HTTPException(
            status_code=400,
            detail="data_json must be a list."
        )

    template_bytes = await (
        template.read()
    )

    result = await process_worker_chunk(

        job_id,

        template_bytes,

        rows,

        node_name,

        coordinator_url.rstrip("/")
    )

    zip_path = Path(
        result["zip_path"]
    )

    if not zip_path.exists():

        raise HTTPException(
            status_code=500,
            detail="Worker ZIP was not created."
        )

    total_elapsed = (
        time.perf_counter()
        -
        started
    )

    log(
        f"{node_name}: "
        f"RETURNING ZIP "
        f"| generation="
        f"{result['elapsed']:.3f}s "
        f"| zip="
        f"{result['zip_elapsed']:.3f}s "
        f"| total="
        f"{total_elapsed:.3f}s"
    )

    log_line()

    return FileResponse(

        path=str(
            zip_path
        ),

        media_type=(
            "application/zip"
        ),

        filename=(
            f"{node_name}.zip"
        ),

        headers={

            "X-Worker-Result":
                json.dumps({

                    "node":
                        node_name,

                    "completed":
                        result[
                            "completed"
                        ],

                    "total":
                        result[
                            "total"
                        ],

                    "generation_time":
                        result[
                            "elapsed"
                        ],

                    "zip_time":
                        result[
                            "zip_elapsed"
                        ],

                    "total_time":
                        total_elapsed
                })
        }
    )


# ============================================================
# STARTUP
# ============================================================

@app.on_event(
    "startup"
)
async def startup():

    log_line()

    log(
        "AUTOGEN NODE STARTED"
    )

    log(
        f"Node: "
        f"{NODE_NAME}"
    )

    log(
        f"Role: "
        f"{NODE_ROLE}"
    )

    log(
        f"Files/sec model: "
        f"{FILES_PER_SECOND}"
    )

    log(
        f"Target seconds: "
        f"{TARGET_SECONDS}"
    )

    log(
        f"Capacity/node: "
        f"{CAPACITY_PER_NODE}"
    )

    if NODE_ROLE == "coordinator":

        log(
            f"Workers configured: "
            f"{len(WORKER_URLS)}"
        )

        for index, url in enumerate(
            WORKER_URLS,
            start=2
        ):

            log(
                f"NODE-{index}: "
                f"{url}"
            )

    else:

        log(
            f"Coordinator: "
            f"{COORDINATOR_URL}"
        )

    log_line()


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    import uvicorn

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=PORT
    )